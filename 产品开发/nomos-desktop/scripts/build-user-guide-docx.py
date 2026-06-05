from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "NOMOS-USER-GUIDE.md"
OUTPUT = ROOT / "release" / "nomos-1.1.0-操作说明书.docx"

INK = RGBColor(24, 45, 39)
GREEN = RGBColor(40, 100, 79)
MUTED = RGBColor(103, 119, 113)


def set_run_font(run, name="Microsoft YaHei", size=10.5, color=INK, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    set_spacing(normal, after=6, line=1.25)

    for style_name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, GREEN),
        ("Heading 2", 13, 14, 7, GREEN),
        ("Heading 3", 11.5, 10, 5, INK),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        set_spacing(style, before=before, after=after, line=1.1)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        set_spacing(style, after=4, line=1.25)


def add_inline_code(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=GREEN, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("nomos Desktop · 本地优先 Agent 项目指挥台")
    set_run_font(header_run, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = document.add_paragraph()
    set_spacing(title, before=22, after=4, line=1.0)
    title_run = title.add_run("nomos Desktop")
    set_run_font(title_run, size=23, color=GREEN, bold=True)

    subtitle = document.add_paragraph()
    set_spacing(subtitle, after=4, line=1.0)
    subtitle_run = subtitle.add_run("1.1.0 操作说明书")
    set_run_font(subtitle_run, size=15, color=INK, bold=True)

    meta = document.add_paragraph()
    set_spacing(meta, after=16, line=1.0)
    meta_run = meta.add_run("适用版本：Windows x64 便携客户端  ·  更新日期：2026-06-05")
    set_run_font(meta_run, size=9.5, color=MUTED)

    lead = document.add_paragraph()
    set_spacing(lead, after=14, line=1.35)
    lead_run = lead.add_run("目标由人确定，过程由 Agent 接力，关键结果由人验收。")
    set_run_font(lead_run, size=12, color=GREEN, bold=True)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()[1:]
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            continue
        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_code(paragraph, number_match.group(1))
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_code(paragraph, line[2:])
            continue
        paragraph = document.add_paragraph()
        set_spacing(paragraph, after=6, line=1.25)
        add_inline_code(paragraph, line)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
